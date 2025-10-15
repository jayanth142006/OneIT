# preprocess.py

import os
from dotenv import load_dotenv
import pdfplumber
from docx import Document as DocxDocument
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.config import Settings
import google.generativeai as genai

# =========== ENV & API KEY ===========
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")
if not api_key:
    raise ValueError("❌ GOOGLE_API_KEY not set in .env")
genai.configure(api_key=api_key)

# =========== Settings ===========
DOCS_FOLDER = "./docs"
CHROMA_DB_PATH = "./chroma_db_placement"
MERGED_TEXT_FILE = "merged_text_output_placement.txt"
COLLECTION_NAME = "placement_documents"

# Embedding model
embedding_model = SentenceTransformer('all-MiniLM-L6-v2')

# =========== PDF Loader with pdfplumber ===========
def load_pdf(filepath):
    """Extract text and tables from PDF with proper formatting"""
    docs = []
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            tables = page.extract_tables()

            table_text = ""
            for table in tables:
                if not table:
                    continue
                
                # Format table as key-value pairs for better understanding
                table_text += "\n--- Table Start ---\n"
                headers = [str(cell).strip() if cell else "" for cell in table[0]]
                
                for row in table[1:]:  # Skip header row
                    for idx, cell in enumerate(row):
                        if idx < len(headers) and headers[idx]:
                            cell_value = str(cell).strip() if cell else "Not specified"
                            table_text += f"{headers[idx]}: {cell_value}\n"
                    table_text += "\n"
                table_text += "--- Table End ---\n\n"

            combined_content = text + "\n" + table_text
            docs.append({
                "content": combined_content,
                "metadata": {"page": i + 1, "source": filepath, "type": "pdf"}
            })
    return docs

# =========== Word Document Loader ===========
def load_docx(filepath):
    """Extract text from Word documents with proper table formatting"""
    doc = DocxDocument(filepath)
    full_text = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # Extract tables with better formatting
    for table_idx, table in enumerate(doc.tables):
        full_text.append(f"\n--- Table {table_idx + 1} Start ---")
        
        # Get headers from first row
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        
        # Process data rows
        for row in table.rows[1:]:
            row_data = []
            for idx, cell in enumerate(row.cells):
                if idx < len(headers) and headers[idx]:
                    cell_value = cell.text.strip() if cell.text.strip() else "Not specified"
                    row_data.append(f"{headers[idx]}: {cell_value}")
            
            if row_data:
                full_text.append("\n".join(row_data))
                full_text.append("")  # Empty line between rows
        
        full_text.append(f"--- Table {table_idx + 1} End ---\n")
    
    content = "\n".join(full_text)
    return [{
        "content": content,
        "metadata": {"source": filepath, "type": "docx"}
    }]

# =========== Load all documents ===========
def load_documents():
    """Load PDF and Word documents from folder"""
    all_docs = []
    
    if not os.path.exists(DOCS_FOLDER):
        os.makedirs(DOCS_FOLDER)
        print(f"📁 Created folder: {DOCS_FOLDER}")
        print(f"⚠️ Please add your PDF and Word documents to {DOCS_FOLDER}")
        return all_docs
    
    for file in os.listdir(DOCS_FOLDER):
        file_path = os.path.join(DOCS_FOLDER, file)
        
        if file.lower().endswith(".pdf"):
            docs = load_pdf(file_path)
            all_docs.extend(docs)
            print(f"✅ Loaded PDF: {file} ({len(docs)} pages)")
            
        elif file.lower().endswith((".docx", ".doc")):
            docs = load_docx(file_path)
            all_docs.extend(docs)
            print(f"✅ Loaded Word: {file}")
    
    print(f"📄 Total documents loaded: {len(all_docs)}")
    return all_docs

# =========== Merge & save ===========
def merge_documents(docs):
    """Merge all document contents and save to text file"""
    merged_text = "\n\n".join([doc["content"] for doc in docs])
    
    with open(MERGED_TEXT_FILE, "w", encoding="utf-8") as f:
        f.write(merged_text)
    
    print(f"📚 Merged text saved to: {MERGED_TEXT_FILE} (length: {len(merged_text)} chars)")
    return merged_text

# =========== Chunk text ===========
def chunk_text(text, chunk_size=1500, chunk_overlap=350):
    """Split text into chunks with overlap"""
    chunks = []
    start = 0
    text_length = len(text)
    
    while start < text_length:
        end = start + chunk_size
        chunk = text[start:end]
        
        # Try to break at sentence boundary
        if end < text_length:
            last_period = chunk.rfind('.')
            last_newline = chunk.rfind('\n')
            break_point = max(last_period, last_newline)
            
            if break_point > chunk_size * 0.5:  # At least 50% of chunk
                chunk = chunk[:break_point + 1]
                end = start + break_point + 1
        
        chunks.append(chunk.strip())
        start = end - chunk_overlap
    
    print(f"✂️ Created {len(chunks)} chunks")
    return chunks

# =========== Create embeddings + ChromaDB ===========
def create_vector_store(chunks):
    """Create embeddings and store in ChromaDB"""
    
    # Initialize ChromaDB client
    client = chromadb.PersistentClient(path=CHROMA_DB_PATH)
    
    # Delete existing collection if it exists
    try:
        client.delete_collection(name=COLLECTION_NAME)
        print("🗑️ Deleted existing collection")
    except:
        pass
    
    # Create new collection
    collection = client.create_collection(
        name=COLLECTION_NAME,
        metadata={"hnsw:space": "cosine"}
    )
    
    print("🔄 Creating embeddings... (this may take a moment)")
    
    # Create embeddings for all chunks
    embeddings = embedding_model.encode(chunks, show_progress_bar=True)
    
    # Prepare data for ChromaDB
    ids = [f"chunk_{i}" for i in range(len(chunks))]
    metadatas = [{"chunk_index": i} for i in range(len(chunks))]
    
    # Add to ChromaDB in batches
    batch_size = 100
    for i in range(0, len(chunks), batch_size):
        batch_chunks = chunks[i:i + batch_size]
        batch_embeddings = embeddings[i:i + batch_size].tolist()
        batch_ids = ids[i:i + batch_size]
        batch_metadatas = metadatas[i:i + batch_size]
        
        collection.add(
            embeddings=batch_embeddings,
            documents=batch_chunks,
            ids=batch_ids,
            metadatas=batch_metadatas
        )
    
    print(f"📦 Vector store created with {len(chunks)} chunks")
    print(f"💾 Saved to: {CHROMA_DB_PATH}")

# =========== Main ===========
if __name__ == "__main__":
    print("=" * 50)
    print("🚀 Starting Preprocessing for Placement Chatbot")
    print("=" * 50)
    
    docs = load_documents()
    
    if not docs:
        print("❌ No documents found. Exiting...")
        exit()
    
    merged_text = merge_documents(docs)
    chunks = chunk_text(merged_text)
    create_vector_store(chunks)
    
    print("=" * 50)
    print("✅ Preprocessing Complete!")
    print("=" * 50)
    print(f"📊 Summary:")
    print(f"   - Documents processed: {len(docs)}")
    print(f"   - Total chunks: {len(chunks)}")
    print(f"   - ChromaDB location: {CHROMA_DB_PATH}")
    print(f"   - Ready to run chatbot with run.py")